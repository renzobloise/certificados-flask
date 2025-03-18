from flask import Flask, render_template, request, send_file
import os
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
import zipfile
import re
import pytz

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
TEMPLATES_FOLDER = 'plantillas'
HISTORIAL = 'historial.csv'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

PLANTILLAS = {
    "no_acepta": "CERTIFICADO NO ACEPTA.docx",
    "no_mediable_generico": "CERTIFICADO NO MEDIABLE  GENERICO.docx",
    "no_mediable_26485": "CERTIFICADO NO MEDIABLE LEY 26485.docx",
    "ignorado_domicilio": "CERTIFICADO IGNORADO DOMICILIO.docx"
}

FIRMANTES = {
    "videla": "DR. JUAN MARTÍN VIDELA - SECRETARIO - CUERPO DE MEDIADORES",
    "bloise": "DR. RENZO A. BLOISE - COORDINADOR GENERAL - CUERPO DE MEDIADORES"
}

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        plantilla = request.form["plantilla"]
        firmante = request.form["firmante"]
        files = request.files.getlist("pdfs")

        generated_files = []
        for file in files:
            path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(path)
            data = extract_data_from_pdf(path)
            docx_path = generate_document(data, plantilla, firmante)
            generated_files.append(docx_path)
            save_to_history(data, plantilla)

        zip_path = os.path.join(OUTPUT_FOLDER, "certificados.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in generated_files:
                zipf.write(file, os.path.basename(file))

        return send_file(zip_path, as_attachment=True)

    return render_template("index.html", plantillas=PLANTILLAS, firmantes=FIRMANTES)

@app.route("/historial")
def historial():
    return send_file(HISTORIAL, as_attachment=True)


def extract_data_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text() for page in doc])
    lines = text.split('\n')

    data = {}
    data['TICKET'] = extract_value(text, r'Trámite:\s*(\d+)')
    data['NOMBRE_SOLICITANTE'] = extract_nombre(lines)
    data['DNI_SOLICITANTE'] = extract_dni(lines)
    data['NOMBRE_PERSONA'] = extract_value(text, r'NOMBRE Y APELLIDO\s*(.*)\n')
    data['DNI_PERSONA'] = extract_value(text, r'NRO\. DE DOCUMENTO\s*(\d+)')

    return {k: v.upper() for k, v in data.items()}


def extract_value(text, pattern):
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1).strip() if match else ""


def extract_nombre(lines):
    for i, line in enumerate(lines):
        if "Apellido y nombre:" in line:
            j = i - 1
            while j >= 0:
                if not re.search(r'\d', lines[j]):
                    return lines[j].strip()
                j -= 1
            return ""
    return ""


def extract_dni(lines):
    for i, line in enumerate(lines):
        if "Documento:" in line and i+1 < len(lines):
            cuil = lines[i+1].strip()
            dni_match = re.search(r'(\d{2})-(\d{8})-(\d{1})', cuil)
            return dni_match.group(2) if dni_match else ""
    return ""


def generate_document(data, plantilla_key, firmante_key):
    template_path = os.path.join(TEMPLATES_FOLDER, PLANTILLAS[plantilla_key])
    doc = Document(template_path)

    argentina_tz = pytz.timezone('America/Argentina/Buenos_Aires')
    fecha_actual = datetime.now(argentina_tz)
    meses_es = {
        "January": "enero", "February": "febrero", "March": "marzo", "April": "abril", "May": "mayo", "June": "junio",
        "July": "julio", "August": "agosto", "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    mes_actual = fecha_actual.strftime("%B")
    mes_es = meses_es.get(mes_actual, mes_actual)
    fecha_formato = fecha_actual.strftime(f"%d días del mes de {mes_es} del año %Y").replace(' 0', ' ')

    # Reemplazo en párrafos normales
    for p in doc.paragraphs:
        for key, value in data.items():
            p.text = p.text.replace(f"{{{{{key}}}}}", value)
        p.text = p.text.replace("{{FECHA_CERTIFICADO}}", fecha_formato)
        p.text = p.text.replace("{{AUTORIDAD_FIRMANTE}}", FIRMANTES[firmante_key])

    # Reemplazo en tablas también
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data.items():
                        p.text = p.text.replace(f"{{{{{key}}}}}", value)
                    p.text = p.text.replace("{{FECHA_CERTIFICADO}}", fecha_formato)
                    p.text = p.text.replace("{{AUTORIDAD_FIRMANTE}}", FIRMANTES[firmante_key])

    output_file = os.path.join(OUTPUT_FOLDER, f"certificado_{data['TICKET']}.docx")
    doc.save(output_file)
    return output_file


def save_to_history(data, plantilla):
    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    registro = f"{fecha},{data['TICKET']},{data['NOMBRE_SOLICITANTE']},{plantilla}\n"
    with open(HISTORIAL, 'a') as f:
        f.write(registro)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
